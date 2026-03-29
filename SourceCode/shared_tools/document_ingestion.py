"""document_ingestion.py — Extract plain text from PDFs, Word docs, plain text, and images (OCR).

Usage:
    from shared_tools.document_ingestion import extract_text, extract_text_from_image, is_document_ext

    text = extract_text(path, mime)            # PDF / DOCX / TXT
    text = extract_text_from_image(path)        # OCR (requires tesseract)
"""
from __future__ import annotations

import logging
from pathlib import Path

from shared_tools.optional_features import feature_warning

LOGGER = logging.getLogger(__name__)

DOCUMENT_MIMES: frozenset[str] = frozenset(
    {
        "application/pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
        "text/plain",
        "text/markdown",
        "text/csv",
    }
)

DOCUMENT_EXTS: frozenset[str] = frozenset({".pdf", ".docx", ".doc", ".txt", ".md", ".csv"})


def is_document_mime(mime: str) -> bool:
    return str(mime).strip().lower() in DOCUMENT_MIMES


def is_document_ext(ext: str) -> bool:
    return str(ext).strip().lower() in DOCUMENT_EXTS


# ---------------------------------------------------------------------------
# Main entry points
# ---------------------------------------------------------------------------

def extract_text(file_path: Path, mime: str = "") -> str:
    """Extract plain text from a document file. Returns empty string on failure."""
    ext = file_path.suffix.lower()
    mime_lower = str(mime).strip().lower()

    try:
        if ext == ".pdf" or "pdf" in mime_lower:
            return _extract_pdf(file_path)
        if ext in {".docx", ".doc"} or "wordprocessingml" in mime_lower or "msword" in mime_lower:
            return _extract_docx(file_path)
        if ext in {".txt", ".md", ".csv"} or mime_lower.startswith("text/"):
            return _extract_plain(file_path)
    except ImportError as exc:
        LOGGER.warning(
            "document_ingestion: optional dependency missing for %s: %s (%s)",
            file_path.name,
            exc,
            feature_warning("document_extraction"),
        )
    except Exception as exc:
        LOGGER.warning("document_ingestion: text extraction failed for %s: %s", file_path.name, exc)

    return ""


def extract_text_from_image(file_path: Path) -> str:
    """Try to extract printed text from an image using OCR (pytesseract + tesseract).
    Returns empty string if tesseract is not installed or OCR yields nothing useful.
    """
    try:
        import pytesseract
        from PIL import Image  # Pillow, bundled with PyMuPDF or installed separately

        img = Image.open(str(file_path))
        text = pytesseract.image_to_string(img).strip()
        return text
    except ImportError:
        LOGGER.info("document_ingestion: OCR unavailable for %s (%s)", file_path.name, feature_warning("image_ocr"))
        return ""
    except Exception as exc:
        LOGGER.debug("document_ingestion: OCR failed for %s: %s", file_path.name, exc)
        return ""


# ---------------------------------------------------------------------------
# Format-specific helpers
# ---------------------------------------------------------------------------

def _extract_pdf(file_path: Path) -> str:
    import fitz  # PyMuPDF

    doc = fitz.open(str(file_path))
    pages: list[str] = []
    for page in doc:
        text = page.get_text()
        if text.strip():
            pages.append(text.strip())
    doc.close()
    return "\n\n".join(pages)


def _extract_docx(file_path: Path) -> str:
    import docx  # python-docx

    document = docx.Document(str(file_path))
    parts: list[str] = []
    for para in document.paragraphs:
        line = para.text.strip()
        if line:
            parts.append(line)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_plain(file_path: Path) -> str:
    for enc in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
        try:
            return file_path.read_text(encoding=enc)
        except (UnicodeDecodeError, ValueError):
            continue
    return file_path.read_bytes().decode("utf-8", errors="replace")
